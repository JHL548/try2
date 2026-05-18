import html
import re
import shutil
import subprocess
import tempfile
import uuid
from pathlib import Path

from bs4 import BeautifulSoup
from bs4.element import Tag
from docx import Document
from fastapi import HTTPException, UploadFile

from app.schemas.documents import NormalizedDocument, RangeMapEntry, TableCellContext


SUPPORTED_WORD_EXTENSIONS = {".doc", ".docx"}
STRUCTURAL_BLOCK_TAGS = {"table", "ul", "ol", "blockquote", "pre"}
TEXT_BLOCK_TAGS = {"p", "h1", "h2", "h3", "li", "td", "th", "div"}
CONTAINER_TAGS = {"body", "main", "section", "article", "div"}
SECTION_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十\d]+[章节部分篇]\s*.+"),
    re.compile(r"^[一二三四五六七八九十]+[、.．]\s*.+"),
    re.compile(r"^\(?[一二三四五六七八九十]\)?[、.．]\s*.+"),
    re.compile(r"^\d+(?:\.\d+)*[、.．]?\s*.+"),
]
LIST_MARKER_PATTERN = re.compile(r"^((?:\d+(?:\.\d+)*|[一二三四五六七八九十]+|\([一二三四五六七八九十\d]+\)|（[一二三四五六七八九十\d]+）|[①②③④⑤⑥⑦⑧⑨⑩])[\s、.．]*)")
SEMANTIC_KEYWORDS = [
    ("priceTable", ("报价", "投标总价", "价格", "费用", "清单")),
    ("deviationTable", ("偏离", "响应情况", "技术偏差", "商务偏差")),
    ("personnel", ("人员", "项目经理", "团队", "配置", "简历")),
    ("performance", ("业绩", "案例", "合同", "类似项目")),
    ("qualification", ("资格", "资质", "证书", "审查", "营业执照")),
    ("schedule", ("进度", "工期", "计划", "里程碑")),
    ("technical", ("技术", "方案", "实施", "服务")),
    ("business", ("商务", "承诺", "条款", "响应")),
    ("projectInfo", ("项目名称", "招标编号", "采购编号", "招标人")),
]


async def upload_to_normalized_document(upload: UploadFile, role: str) -> NormalizedDocument:
    suffix = Path(upload.filename or "document.docx").suffix.lower()
    if suffix not in SUPPORTED_WORD_EXTENSIONS:
        raise HTTPException(status_code=400, detail="当前仅支持 DOC 或 DOCX 文档")

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / f"input{suffix or '.docx'}"
        temp_path.write_bytes(await upload.read())
        html_content = convert_file_to_html(temp_path)

    return html_to_normalized_document(
        html_content,
        file_name=upload.filename or "document.docx",
        role=role,
    )


def convert_file_to_html(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix in SUPPORTED_WORD_EXTENSIONS:
        libreoffice_html = convert_with_libreoffice(file_path, "html")
        if libreoffice_html:
            return libreoffice_html.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".docx":
            return docx_to_basic_html(file_path)

    raise ValueError("Unsupported file extension")


def convert_with_libreoffice(file_path: Path, target_format: str) -> Path | None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None

    output_dir = file_path.parent / "converted"
    output_dir.mkdir(exist_ok=True)

    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                target_format,
                "--outdir",
                str(output_dir),
                str(file_path),
            ],
            check=True,
            capture_output=True,
            timeout=60,
        )
    except (subprocess.SubprocessError, OSError):
        return None

    converted_files = list(output_dir.glob(f"*.{target_format.split(':')[0]}"))
    return converted_files[0] if converted_files else None


def docx_to_basic_html(file_path: Path) -> str:
    document = Document(file_path)
    paragraphs: list[str] = []

    for index, paragraph in enumerate(document.paragraphs):
        text = html.escape(paragraph.text)
        if not text:
            paragraphs.append(f'<p data-dupdoc-block="p-{index}"><br></p>')
            continue
        paragraphs.append(f'<p data-dupdoc-block="p-{index}">{text}</p>')

    return "\n".join(paragraphs) or '<p data-dupdoc-block="p-0"></p>'


def html_to_normalized_document(html_content: str, file_name: str, role: str) -> NormalizedDocument:
    soup = BeautifulSoup(html_content, "html.parser")
    body = soup.body or soup
    normalized_parts: list[str] = []
    plain_text_parts: list[str] = []
    range_map: list[RangeMapEntry] = []
    cursor = 0
    section_path: list[str] = []

    block_nodes = collect_normalized_blocks(body)
    if not block_nodes:
        block_nodes = [body]

    for index, node in enumerate(block_nodes):
        if not isinstance(node, Tag):
            continue

        text = node.get_text(" ", strip=True)
        if not text and node.name != "br":
            continue

        block_id = node.get("data-dupdoc-block") or f"b-{index}"
        node["data-dupdoc-block"] = block_id
        list_marker = extract_list_marker(text)
        if is_section_heading(node, text):
            section_path = update_section_path(section_path, text)

        table_context = build_table_context(node)
        region = "table" if table_context or node.name == "table" else "body"
        semantic_type = infer_semantic_type(text, section_path)

        annotate_structural_node(node, block_id)
        node["data-dupdoc-region"] = region
        if section_path:
            node["data-dupdoc-section"] = " / ".join(section_path)
        if semantic_type:
            node["data-dupdoc-semantic-type"] = semantic_type

        rendered_html = str(node)

        if normalized_parts and text:
            plain_text_parts.append("\n")
            cursor += 1

        start = cursor
        plain_text_parts.append(text)
        cursor += len(text)
        end = cursor

        normalized_parts.append(rendered_html)
        range_map.append(
            RangeMapEntry(
                blockId=block_id,
                textStart=start,
                textEnd=end,
                selector=f'[data-dupdoc-block="{block_id}"]',
                sectionPath=section_path or None,
                region=region,
                listMarker=list_marker,
                tableContext=table_context,
                semanticType=semantic_type,
            )
        )

    plain_text = "".join(plain_text_parts)

    return NormalizedDocument(
        documentId=str(uuid.uuid4()),
        role=role,  # type: ignore[arg-type]
        fileName=file_name,
        html="\n".join(normalized_parts) or '<p data-dupdoc-block="b-0"></p>',
        plainText=plain_text,
        rangeMap=range_map,
        meta={
            "sourceFileName": file_name,
            "converter": "libreoffice-or-python-fallback",
            "formatLossExpected": True,
        },
    )


def collect_normalized_blocks(root: Tag) -> list[Tag]:
    blocks: list[Tag] = []

    for child in root.children:
        if not isinstance(child, Tag):
            continue

        if child.name in STRUCTURAL_BLOCK_TAGS:
            blocks.append(child)
            continue

        if child.name in TEXT_BLOCK_TAGS and not has_structural_child(child):
            blocks.append(child)
            continue

        if child.name in CONTAINER_TAGS:
            nested_blocks = collect_normalized_blocks(child)
            if nested_blocks:
                blocks.extend(nested_blocks)
            elif child.get_text(" ", strip=True):
                blocks.append(child)
            continue

        nested_blocks = collect_normalized_blocks(child)
        if nested_blocks:
            blocks.extend(nested_blocks)
        elif child.get_text(" ", strip=True):
            blocks.append(child)

    return blocks


def has_structural_child(node: Tag) -> bool:
    return node.find(["table", "ul", "ol", "blockquote", "pre"]) is not None


def annotate_structural_node(node: Tag, block_id: str) -> None:
    node["data-dupdoc-block"] = block_id
    if node.name == "table":
        ensure_table_id(node)
        for cell in node.find_all(["td", "th"]):
            cell["data-dupdoc-block"] = block_id
    if node.name in {"ul", "ol"}:
        for item in node.find_all("li"):
            item["data-dupdoc-block"] = block_id


def is_section_heading(node: Tag, text: str) -> bool:
    if node.name in {"h1", "h2", "h3"}:
        return True
    return any(pattern.match(text) for pattern in SECTION_PATTERNS)


def update_section_path(current_path: list[str], heading: str) -> list[str]:
    level = infer_heading_level(heading)
    next_path = current_path[: max(level - 1, 0)]
    next_path.append(heading[:80])
    return next_path


def infer_heading_level(heading: str) -> int:
    if heading.startswith("第"):
        return 1
    if re.match(r"^[一二三四五六七八九十]+[、.．]", heading):
        return 2
    if re.match(r"^\(?[一二三四五六七八九十]\)?[、.．]", heading):
        return 3
    dotted = re.match(r"^(\d+(?:\.\d+)*)", heading)
    if dotted:
        return min(dotted.group(1).count(".") + 1, 4)
    return 2


def extract_list_marker(text: str) -> str | None:
    match = LIST_MARKER_PATTERN.match(text)
    return match.group(1).strip() if match else None


def infer_semantic_type(text: str, section_path: list[str]) -> str | None:
    context = " ".join([*section_path, text])
    for semantic_type, keywords in SEMANTIC_KEYWORDS:
        if any(keyword in context for keyword in keywords):
            return semantic_type
    return None


def build_table_context(node: Tag) -> TableCellContext | None:
    if node.name == "table":
        first_data_cell = find_first_table_cell(node)
        if first_data_cell:
            return build_table_context(first_data_cell)
        return None

    if node.name not in {"td", "th"}:
        return None

    table = node.find_parent("table")
    row = node.find_parent("tr")
    if not table or not row:
        return None

    rows = [item for item in table.find_all("tr", recursive=False)]
    if not rows:
        rows = table.find_all("tr")
    cells = [item for item in row.find_all(["td", "th"], recursive=False)]

    try:
        row_index = rows.index(row)
        cell_index = cells.index(node)
    except ValueError:
        return None

    table_id = ensure_table_id(table)
    header_text = get_table_header_text(rows, cell_index)
    row_header_text = cells[0].get_text(" ", strip=True) if cells and cell_index > 0 else None

    return TableCellContext(
        tableId=table_id,
        rowIndex=row_index,
        cellIndex=cell_index,
        headerText=header_text,
        rowHeaderText=row_header_text,
    )


def ensure_table_id(table: Tag) -> str:
    table_id = table.get("data-dupdoc-table") or f"table-{abs(hash(str(table)[:120]))}"
    table["data-dupdoc-table"] = table_id
    return str(table_id)


def find_first_table_cell(table: Tag) -> Tag | None:
    rows = table.find_all("tr")
    if len(rows) > 1:
        cell = rows[1].find(["td", "th"])
        if isinstance(cell, Tag):
            return cell

    cell = table.find(["td", "th"])
    return cell if isinstance(cell, Tag) else None


def get_table_header_text(rows: list[Tag], cell_index: int) -> str | None:
    if not rows:
        return None
    header_cells = rows[0].find_all(["td", "th"], recursive=False)
    if cell_index >= len(header_cells):
        return None
    return header_cells[cell_index].get_text(" ", strip=True) or None


def html_to_docx_file(html_content: str, file_name: str) -> Path:
    temp_dir = Path(tempfile.mkdtemp())
    html_path = temp_dir / "edited.html"
    html_path.write_text(wrap_html_document(html_content), encoding="utf-8")

    converted = convert_with_libreoffice(html_path, "docx")
    if converted:
        output_path = temp_dir / ensure_docx_name(file_name)
        converted.replace(output_path)
        return output_path

    output_path = temp_dir / ensure_docx_name(file_name)
    document = Document()
    soup = BeautifulSoup(html_content, "html.parser")
    for block in soup.find_all(["p", "h1", "h2", "h3", "li", "div", "table", "blockquote", "pre"]):
        document.add_paragraph(block.get_text(" ", strip=True))
    document.save(output_path)
    return output_path


def wrap_html_document(html_content: str) -> str:
    return f"""<!doctype html>
<html>
  <head>
    <meta charset="utf-8">
  </head>
  <body>{html_content}</body>
</html>
"""


def ensure_docx_name(file_name: str) -> str:
    path = Path(file_name)
    return f"{path.stem or 'edited-document'}.docx"
